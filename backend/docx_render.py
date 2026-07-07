"""
Render an artifact definition into a Word .docx.

Two modes:
- Blank template mode: produce a reference-manual-depth template with empty tables
  and fillable sections. Used to ship the pack of 21 templates.
- Project mode: fill fields, tables, and add a project header with all sections
  auto-populated where data exists.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches


# Design tokens (neutral corporate with a Hydra-teal accent)
COLOR_TEXT = RGBColor(0x28, 0x25, 0x1D)
COLOR_MUTED = RGBColor(0x7A, 0x79, 0x74)
COLOR_ACCENT = RGBColor(0x01, 0x69, 0x6F)
COLOR_HEADER_BG = "01696F"
COLOR_ROW_ALT_BG = "F1EFEA"
COLOR_TABLE_BORDER = "D4D1CA"


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), COLOR_TABLE_BORDER)
        borders.append(b)
    tbl_pr.append(borders)


def _apply_run_style(run, *, bold=False, italic=False, size=10.5, color=COLOR_TEXT, font="Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 0:
        _apply_run_style(run, bold=True, size=22, color=COLOR_ACCENT, font="Calibri")
    elif level == 1:
        _apply_run_style(run, bold=True, size=16, color=COLOR_ACCENT, font="Calibri")
    elif level == 2:
        _apply_run_style(run, bold=True, size=13, color=COLOR_TEXT, font="Calibri")
    else:
        _apply_run_style(run, bold=True, size=11, color=COLOR_TEXT, font="Calibri")


def _add_paragraph(doc: Document, text: str, *, bold=False, italic=False, size=10.5, color=COLOR_TEXT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _apply_run_style(run, bold=bold, italic=italic, size=size, color=color)


def _add_bullet(doc: Document, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_run_style(run, size=size)


def _add_numbered(doc: Document, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_run_style(run, size=size)


def _add_callout(doc: Document, label: str, text: str) -> None:
    # A single-cell shaded table used as a callout box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    _set_table_borders(tbl)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, COLOR_ROW_ALT_BG)
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}  ")
    _apply_run_style(run, bold=True, size=10.5, color=COLOR_ACCENT)
    run = p.add_run(text)
    _apply_run_style(run, size=10.5)
    doc.add_paragraph()  # spacer


def _add_table(doc: Document, columns: list[dict[str, str]], rows: list[list[str]] | None = None,
               min_blank_rows: int = 3) -> None:
    n_cols = len(columns)
    n_rows = 1 + max(len(rows or []), min_blank_rows)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.autofit = True
    _set_table_borders(tbl)

    # Header row
    for i, col in enumerate(columns):
        cell = tbl.rows[0].cells[i]
        _set_cell_shading(cell, COLOR_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(col["label"])
        _apply_run_style(run, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    data_rows = rows or []
    for r_idx in range(1, n_rows):
        alt = (r_idx % 2 == 0)
        for c_idx, col in enumerate(columns):
            cell = tbl.rows[r_idx].cells[c_idx]
            if alt:
                _set_cell_shading(cell, COLOR_ROW_ALT_BG)
            val = ""
            if r_idx - 1 < len(data_rows):
                row = data_rows[r_idx - 1]
                if c_idx < len(row):
                    val = str(row[c_idx]) if row[c_idx] is not None else ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _apply_run_style(run, size=10)


def _add_field_form(doc: Document, fields: list[dict[str, Any]], values: dict[str, Any] | None = None) -> None:
    """Render fillable fields as a 2-column table (label | value)."""
    if not fields:
        return
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.autofit = True
    _set_table_borders(tbl)
    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)
    for i, f in enumerate(fields):
        label_cell = tbl.rows[i].cells[0]
        _set_cell_shading(label_cell, COLOR_ROW_ALT_BG)
        p = label_cell.paragraphs[0]
        run = p.add_run(f["label"])
        _apply_run_style(run, bold=True, size=10)

        value_cell = tbl.rows[i].cells[1]
        p = value_cell.paragraphs[0]
        val = ""
        if values and f["key"] in values and values[f["key"]] not in (None, ""):
            v = values[f["key"]]
            if isinstance(v, list):
                val = ", ".join(str(x) for x in v)
            else:
                val = str(v)
        run = p.add_run(val)
        _apply_run_style(run, size=10)
        if f.get("help"):
            p2 = value_cell.add_paragraph()
            run = p2.add_run(f["help"])
            _apply_run_style(run, italic=True, size=9, color=COLOR_MUTED)


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()


def _add_footer(doc: Document, artifact: dict[str, Any], project_name: str | None) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt = f"Artifact {artifact['number']:02d} · {artifact['title']}"
    if project_name:
        txt = f"{project_name} · {txt}"
    txt += f" · Generated {datetime.now().strftime('%Y-%m-%d')}"
    run = p.add_run(txt)
    _apply_run_style(run, size=9, color=COLOR_MUTED)


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def render_artifact(
    artifact: dict[str, Any],
    *,
    project_name: str | None = None,
    domain: str = "generic",
    field_values: dict[str, Any] | None = None,
    table_rows: dict[str, list[list[str]]] | None = None,
) -> bytes:
    """Render a single artifact to DOCX bytes."""
    doc = Document()
    _configure_page(doc)
    _add_footer(doc, artifact, project_name)

    # ---- Cover / header ----
    _add_heading(doc, f"Artifact {artifact['number']:02d}", level=2)
    _add_heading(doc, artifact["title"], level=0)
    _add_paragraph(doc, "AI Cost Architect — Principal Consultant — Reference Manual Template",
                   italic=True, color=COLOR_MUTED, size=11)
    if project_name:
        _add_paragraph(doc, f"Project: {project_name}    Domain overlay: {domain}",
                       bold=True, size=11, color=COLOR_ACCENT)
    doc.add_paragraph()

    # ---- Role context ----
    _add_heading(doc, "Role context — why the Principal AI Solutions Consultant owns this", level=1)
    _add_paragraph(doc, artifact["role_context"])

    # ---- Purpose / when / inputs ----
    _add_heading(doc, "Purpose", level=1)
    _add_paragraph(doc, artifact["purpose"])
    _add_heading(doc, "When to use", level=1)
    _add_paragraph(doc, artifact["when_to_use"])
    _add_heading(doc, "Inputs needed", level=1)
    for x in artifact["inputs_needed"]:
        _add_bullet(doc, x)

    # ---- Concepts ----
    learning = artifact["learning"]
    if learning.get("concepts"):
        _add_heading(doc, "Key concepts (glossary)", level=1)
        _add_paragraph(doc,
            "If you have never done AI cost management before, read this section carefully. "
            "Each concept is used elsewhere in this artifact and across the pack.",
            italic=True, color=COLOR_MUTED)
        for c in learning["concepts"]:
            _add_heading(doc, c["term"], level=3)
            _add_paragraph(doc, "Definition: " + c["definition"])
            _add_paragraph(doc, "Why it matters: " + c["why_it_matters"])

    # ---- How to use ----
    if learning.get("how_to_use"):
        _add_heading(doc, "How to use this artifact (step-by-step)", level=1)
        for step in learning["how_to_use"]:
            _add_numbered(doc, step)

    # ---- Worked example ----
    if learning.get("worked_example"):
        _add_heading(doc, "Worked example", level=1)
        _add_callout(doc, "Example:", learning["worked_example"])

    # ---- Pitfalls ----
    if learning.get("pitfalls"):
        _add_heading(doc, "Common pitfalls (and how to avoid them)", level=1)
        for p in learning["pitfalls"]:
            _add_bullet(doc, p)

    # ---- Decision tree ----
    if learning.get("decision_tree"):
        _add_heading(doc, "Decision tree", level=1)
        for d in learning["decision_tree"]:
            _add_bullet(doc, d)

    # ---- Domain overlays ----
    if learning.get("domain_overlays"):
        _add_heading(doc, "Domain overlays", level=1)
        _add_paragraph(doc,
            "This template is domain-generic. Add the following considerations depending on the "
            "domain you are operating in. For any regulated domain, the domain overlay must be "
            "socialized with compliance before publishing the artifact.",
            italic=True, color=COLOR_MUTED)
        overlays = learning["domain_overlays"]
        # Highlight the selected overlay first
        if domain in overlays:
            _add_heading(doc, f"Selected: {domain}", level=3)
            _add_callout(doc, "Overlay:", overlays[domain])
        _add_heading(doc, "All overlays for reference", level=3)
        for k, v in overlays.items():
            _add_paragraph(doc, f"{k}: {v}")

    _add_page_break(doc)

    # ---- Fillable form ----
    if artifact.get("fields"):
        _add_heading(doc, "Fillable form", level=1)
        _add_paragraph(doc,
            "Complete the fields below. When rendered from the AI Cost Architect Portfolio "
            "application, these values are populated automatically from the project record.",
            italic=True, color=COLOR_MUTED)
        _add_field_form(doc, artifact["fields"], values=field_values)

    # ---- Tables ----
    if artifact.get("tables"):
        for table_def in artifact["tables"]:
            _add_heading(doc, table_def["label"], level=1)
            rows = None
            if table_rows and table_def["key"] in table_rows:
                data = table_rows[table_def["key"]]
                # data is list of dicts; convert to rows in column order
                rows = []
                for r in data:
                    rows.append([r.get(c["key"], "") for c in table_def["columns"]])
            _add_table(doc, table_def["columns"], rows=rows, min_blank_rows=4)

    # ---- Notes ----
    _add_heading(doc, "Notes", level=1)
    _add_paragraph(doc, "Use this space for open questions, follow-ups, and links to related artifacts.")
    for _ in range(3):
        doc.add_paragraph()

    # Return bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_pack_zip(artifacts: list[dict[str, Any]], out_dir: str) -> list[str]:
    """Render all artifacts as blank templates to disk."""
    import os
    os.makedirs(out_dir, exist_ok=True)
    files = []
    for a in artifacts:
        data = render_artifact(a)
        # sanitize filename
        safe_title = a["short_title"].replace("/", "-").replace(" ", "_")
        fn = f"Artifact_{a['number']:02d}_{safe_title}.docx"
        path = os.path.join(out_dir, fn)
        with open(path, "wb") as f:
            f.write(data)
        files.append(path)
    return files
